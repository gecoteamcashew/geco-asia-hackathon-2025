# team_cashew_app.py
# Nibbles (customer) + Clarity (staff)
# - Robust CSV ingestion (recursive finder, delimiter/encoding inference, header/value cleanup)
# - Product index from SKU Master, Sales Descriptions, and optional Catalogue_Flat
# - Customer replies: friendly, data-only (FAQ / CSVs). No AI facts.
# - Staff analytics: require "Staff query:"; default to PRODUCT NAMES (not SKUs).
#   Only show SKUs if explicitly asked (e.g., "with SKU", "by SKU", "show SKUs", "product code").
# - Multilingual I/O focused on SEA + CJK: script/lexicon detection, translation-only via GPT-OSS.
# - Precise small-talk handler for English + SEA/CJK greetings (no short-text heuristic).
# - Debug endpoints: /debug/sources, /debug/search, /health

import os, sys, re, subprocess, logging
from pathlib import Path
from typing import Dict, Any, Optional, Tuple, List
from datetime import datetime, timedelta
from string import Template

# -----------------------------
# Logging
# -----------------------------
logging.basicConfig(level=logging.INFO, format="%(message)s")
log = logging.getLogger("team_cashew_bot")

# -----------------------------
# Auto-install dependencies
# -----------------------------
REQUIRED = [
    "flask","pandas","python-dateutil","waitress",
    "fuzzywuzzy","python-levenshtein","python-docx","openpyxl","unidecode","requests"
]
def ensure_deps():
    import importlib
    missing=[]
    for pkg in REQUIRED:
        mod = "Levenshtein" if pkg=="python-levenshtein" else \
              "dateutil" if pkg=="python-dateutil" else \
              "docx" if pkg=="python-docx" else pkg
        try: importlib.import_module(mod)
        except ImportError: missing.append(pkg)
    if missing:
        print(f"Installing missing packages: {missing}")
        subprocess.check_call([sys.executable,"-m","pip","install",*missing])
ensure_deps()

# -----------------------------
# Imports (post-install)
# -----------------------------
import pandas as pd
from flask import Flask, request, jsonify, render_template_string
from fuzzywuzzy import fuzz
from docx import Document
from unidecode import unidecode
import requests

# -----------------------------
# Config
# -----------------------------
ASSISTANT_NAME_CUSTOMER = "Nibbles"
ASSISTANT_NAME_STAFF    = "Clarity"

BITDEER_API_KEY=os.getenv("BITDEER_API_KEY","").strip()
BITDEER_URL="https://api-inference.bitdeer.ai/v1/chat/completions"
BITDEER_MODEL="openai/gpt-oss-120b"

DEFAULT_DATA_DIR = Path(os.getenv("DATA_DIR", r"C:\Team_Cashew_Synthetic_Data"))
DATA_DIR = DEFAULT_DATA_DIR

FILE_HINTS = {
    "faq": ["faq_list","master_faq_list","faq"],
    "catalogue": ["camel product catalogue","catalogue","catalog"],
    "sku_master": ["sku_master","sku master","sku"],
    "sales_transactions": ["sales_transactions","sales transactions","transactions","sales"],
    "ecommerce_purchases": ["ecommerce_purchases","ecommerce purchases","purchases"],
    "traffic_acquisition": ["traffic_acquisition","traffic acquisition","acquisition","traffic"],
    "events": ["events","event"],
    "customers": ["customer_table","customer table","customers"]
}

# -----------------------------
# Translation helpers (translation-only; no added facts)
# -----------------------------
def _call_gpt_translate(system_prompt: str, user_text: str, max_tokens=300) -> str:
    if not BITDEER_API_KEY:
        return ""  # translation disabled -> caller will fallback to original
    try:
        r = requests.post(
            BITDEER_URL,
            headers={"Authorization": f"Bearer {BITDEER_API_KEY}", "Content-Type": "application/json"},
            json={
                "model": BITDEER_MODEL,
                "messages": [
                    {"role":"system","content":system_prompt},
                    {"role":"user","content":user_text}
                ],
                "max_tokens": max_tokens,
                "temperature": 0.0,
                "top_p": 1,
                "frequency_penalty": 0,
                "presence_penalty": 0,
                "stream": False
            },
            timeout=20
        )
        txt = (r.json().get("choices",[{}])[0].get("message",{}) or {}).get("content","")
        txt = re.sub(r"[*_`>#]", " ", txt or "")
        txt = re.sub(r"\s+\n", "\n", txt)
        return txt.strip()
    except Exception:
        return ""

# ---------- SEA/CJK language detection (script + lexicon) ----------
JP_HIRA = re.compile(r'[\u3040-\u309F]')          # Hiragana
JP_KATA = re.compile(r'[\u30A0-\u30FF]')          # Katakana
CN_HAN  = re.compile(r'[\u4E00-\u9FFF]')          # CJK Unified (Han)
KO_HANG = re.compile(r'[\uAC00-\uD7AF]')          # Hangul
TH_TH   = re.compile(r'[\u0E00-\u0E7F]')          # Thai
LAO_LA  = re.compile(r'[\u0E80-\u0EFF]')          # Lao
KM_KH   = re.compile(r'[\u1780-\u17FF]')          # Khmer
MY_MM   = re.compile(r'[\u1000-\u109F]')          # Myanmar
VI_DIAC = re.compile(r'[ăâđêôơưĂÂĐÊÔƠƯ]')         # Vietnamese precomposed letters

MS_ID_TOKENS = {
    "ms": {"apa", "boleh", "pesan", "penghantaran", "terima kasih", "selamat", "maklumat", "barangan", "kacang"},
    "id": {"apa", "bisa", "pesan", "pengiriman", "terima kasih", "selamat", "informasi", "barang", "kacang"},
}
TL_TOKENS = {"kumusta", "magkano", "salamat", "po", "paano", "meron", "wala", "balik", "utos", "paki"}

def detect_lang_safe(text: str) -> str:
    t = text.strip()
    if not t:
        return "en"
    if JP_HIRA.search(t) or JP_KATA.search(t): return "ja"
    if KO_HANG.search(t):                      return "ko"
    if CN_HAN.search(t):                       return "zh"
    if TH_TH.search(t):                        return "th"
    if LAO_LA.search(t):                       return "lo"
    if KM_KH.search(t):                        return "km"
    if MY_MM.search(t):                        return "my"
    if VI_DIAC.search(t):                      return "vi"
    tl = t.lower()
    if any(tok in tl for tok in TL_TOKENS):           return "tl"
    if any(tok in tl for tok in MS_ID_TOKENS["ms"]):  return "ms"
    if any(tok in tl for tok in MS_ID_TOKENS["id"]):  return "id"
    return "en"

def needs_translation(lang: str) -> bool:
    return lang != "en"

def to_english_keywords(text: str, user_lang: str) -> str:
    if not needs_translation(user_lang) or not text.strip():
        return text
    sys_prompt = "Translate the user's message into concise English keywords for search. Do not add or change meaning. Plain text only."
    out = _call_gpt_translate(sys_prompt, text, max_tokens=80)
    return out or text

def from_english_plain(text: str, user_lang: str) -> str:
    if not needs_translation(user_lang) or not text.strip():
        return text
    sys_prompt = (
        f"Translate into {user_lang} in plain text. Preserve product names and numbers exactly. "
        "Do not add new information or change facts. No markdown."
    )
    out = _call_gpt_translate(sys_prompt, text, max_tokens=400)
    return out or text

# -----------------------------
# Friendly tone helpers (customer only)
# -----------------------------
FRIENDLY_GREETS = [
    "Hi there! Happy to help.",
    "Hello! What can I do for you today?",
    "Hey! I’m here for your snack questions.",
]
FRIENDLY_OK = [
    "Sure — here’s what I found.",
    "Got it — here’s what I can share.",
    "Thanks for asking — here’s the info.",
]
FRIENDLY_SORRY = [
    "Sorry — I couldn’t find a perfect match.",
    "Hmm, I couldn’t find that exactly.",
    "I didn’t see an exact match yet.",
]
def _pick(items): return items[0]  # deterministic

# -----------------------------
# Precise small-talk (English + SEA/CJK), NO short-text fallback
# -----------------------------
SMALLTALK_EN = re.compile(
    r"^(?:hi|hello|hey|good (?:morning|afternoon|evening)|thanks|thank you|bye|goodbye)[.! ]*$",
    re.I
)
GREET_SEA = re.compile(
    r"(?:"
    r"你好|您好|嗨|哈囉|こんにちは|こんばんは|おはよう|안녕하세요|안녕|"
    r"สวัสดี|ສະບາຍດີ|សួស្តី|မင်္ဂလာပါ|"
    r"xin chào|chào|cảm ơn|"
    r"selamat pagi|selamat siang|selamat sore|selamat malam|terima kasih|apa khabar|apa kabar|"
    r"kumusta|magandang umaga|magandang gabi|salamat"
    r")$",
    re.I
)
INTENT_WORDS = re.compile(
    r"\b(order|refund|return|ship|shipping|deliver|delivery|arrive|arrival|eta|status|track|tracking|allergen|ingredient|have|do you have|find|search)\b",
    re.I
)
INTENT_SEA = re.compile(
    r"(?:"
    # Chinese
    r"有.*吗|有吗|有没有|想买|购买|下单|退款|退货|配送|运费|寄送|发货|送货|多久|几天|到货|查询|追踪"
    r"|"
    # Japanese
    r"ありますか|購入|買えますか|注文|返品|返金|配達|配送|送料|どのくらい|到着|追跡|トラッキング"
    r"|"
    # Korean
    r"있나요|있습니까|구매|주문|환불|반품|배송|얼마|며칠|도착|추적"
    r"|"
    # Vietnamese
    r"có.*không|mua|đặt hàng|hoàn tiền|trả hàng|vận chuyển|giao hàng|bao lâu|mất mấy ngày|đến khi nào|theo dõi|giá"
    r"|"
    # Malay / Indonesian
    r"ada|punya|pesan|tempah|penghantaran|pengiriman|bayaran balik|refund|pulangan|berapa hari|berapa lama|hantar|sampai bila|jejak|penjejakan|harga"
    r"|"
    # Tagalog / Filipino
    r"meron|mayroon|paano|magkano|mag-order|umorder|i-refund|refund|ibalik|padala|ilang araw|gaano katagal|darating|i-track|subaybayan|presyo"
    r")",
    re.I
)
def is_smalltalk(message: str) -> bool:
    s = message.strip()
    sl = s.lower()
    if INTENT_WORDS.search(sl) or INTENT_SEA.search(s):
        return False
    if SMALLTALK_EN.match(sl):
        return True
    lang = detect_lang_safe(s)
    if lang != "en" and GREET_SEA.search(sl):
        return True
    return False

# -----------------------------
# Date helpers
# -----------------------------
def parse_date_range(text: str) -> Tuple[Optional[datetime], Optional[datetime]]:
    q = text.lower()
    now = datetime.now()
    if "last week" in q:
        return now - timedelta(days=7), now
    if "last month" in q:
        first = now.replace(day=1)
        last = first - timedelta(days=1)
        return last.replace(day=1), last
    m_q = re.search(r"\b(q[1-4])\s*(20\d{2})\b", q, re.I)
    if m_q:
        qtr = m_q.group(1).upper(); y = int(m_q.group(2))
        return {
            "Q1": (datetime(y,1,1),   datetime(y,3,31,23,59,59)),
            "Q2": (datetime(y,4,1),   datetime(y,6,30,23,59,59)),
            "Q3": (datetime(y,7,1),   datetime(y,9,30,23,59,59)),
            "Q4": (datetime(y,10,1),  datetime(y,12,31,23,59,59)),
        }[qtr]
    m_y = re.search(r"\b(20\d{2})\b", q)
    if m_y:
        y = int(m_y.group(1))
        return datetime(y,1,1), datetime(y,12,31,23,59,59)
    return None, None

def filter_by_date(df: pd.DataFrame, start, end):
    if df.empty or (start is None and end is None):
        return df
    date_cols = [c for c in df.columns if re.search(r"(date|time)", str(c), re.I)]
    if not date_cols:
        return df
    best_col, best_series, best_valid = None, None, -1
    for col in date_cols:
        raw = df[col]
        candidates = [
            pd.to_datetime(raw, errors="coerce", infer_datetime_format=True),
            pd.to_datetime(raw, errors="coerce", dayfirst=True, infer_datetime_format=True),
        ]
        for ser in candidates:
            valid = ser.notna().sum()
            if valid > best_valid:
                best_col, best_series, best_valid = col, ser, valid
    if best_valid <= 0:
        return df
    out = df.copy()
    out[best_col] = best_series
    before = len(out)
    if start is not None: out = out[out[best_col] >= start]
    if end   is not None: out = out[out[best_col] <= end]
    after = len(out)
    log.info(f"[date-filter] Using '{best_col}' valid={best_valid} rows {before}->{after}")
    return out

# -----------------------------
# General utils
# -----------------------------
def _norm(s: str) -> str:
    return re.sub(r'[^a-z0-9]+', '', str(s).lower())

# -----------------------------
# Data Manager
# -----------------------------
class DataManager:
    def __init__(self, data_dir: Path):
        self.dir = Path(data_dir)
        self.tables: Dict[str, pd.DataFrame] = {}
        self.catalogue = pd.DataFrame()
        self.catalogue_flat = pd.DataFrame()
        self.product_index: List[str] = []
        self._load()

    def _find_file(self, hints, exts):
        hints_n = [_norm(h) for h in hints]
        candidates = []
        for f in self.dir.rglob("*"):
            if f.is_file() and f.suffix.lower() in exts:
                fn = _norm(f.name)
                best = 0
                for hn in hints_n:
                    pos = fn.find(hn)
                    if pos >= 0:
                        best = max(best, 100 - pos)
                if best > 0:
                    candidates.append((best, f))
        if not candidates:
            log.info(f"[find] No match for hints={hints} in {self.dir}")
            return None
        candidates.sort(key=lambda x: x[0], reverse=True)
        chosen = candidates[0][1]
        alts = ", ".join(str(p) for _, p in candidates[:5])
        log.info(f"[find] hints={hints} -> {chosen}  (alts: {alts})")
        return chosen

    def _safe_csv(self, p: Path) -> pd.DataFrame:
        for enc in ("utf-8","utf-8-sig","latin-1"):
            try:
                df = pd.read_csv(p, encoding=enc, sep=None, engine="python")
                break
            except Exception:
                df = pd.DataFrame()
        if df.empty:
            log.info(f"[csv] Failed to read: {p}")
            return df
        df.columns = [str(c).replace("\u00A0"," ").strip() for c in df.columns]
        for c in df.columns:
            if df[c].dtype == object:
                df[c] = df[c].astype(str).str.replace("\u00A0"," ", regex=False).str.replace(r"\s+"," ", regex=True).str.strip()
        return df

    def _safe_excel(self, p: Path):
        try:
            df = pd.read_excel(p)
            df.columns = [str(c).replace("\u00A0"," ").strip() for c in df.columns]
            for c in df.columns:
                if df[c].dtype == object:
                    df[c] = df[c].astype(str).str.replace("\u00A0"," ", regex=False).str.replace(r"\s+"," ", regex=True).str.strip()
            return df
        except Exception:
            log.info(f"[xlsx] Failed to read: {p}")
            return pd.DataFrame()

    def _parse_catalogue_docx(self, p: Optional[Path]):
        if not p: return pd.DataFrame()
        try:
            doc = Document(str(p))
        except Exception:
            return pd.DataFrame()
        lines = [x.text.strip() for x in doc.paragraphs if x.text.strip()]
        return pd.DataFrame({"name": list(set(lines))})

    def _flatten_catalogue_to_csv(self):
        if self.catalogue.empty:
            return
        col = next((c for c in self.catalogue.columns if "name" in c.lower()), self.catalogue.columns[0])
        flat = self.catalogue[[col]].rename(columns={col: "Product_Name"})
        flat = flat[flat["Product_Name"].astype(str).str.len() >= 3].drop_duplicates()
        self.catalogue_flat = flat.reset_index(drop=True)
        out = self.dir / "Catalogue_Flat.csv"
        try:
            self.catalogue_flat.to_csv(out, index=False, encoding="utf-8-sig")
            log.info(f"[catalogue] Flattened DOCX -> {out} ({len(self.catalogue_flat)} rows)")
        except Exception as e:
            log.info(f"[catalogue] Flatten failed: {e}")

    def _normalize_faq(self, df: pd.DataFrame):
        if df.empty: return df
        cols = [str(c).strip().lower() for c in df.columns]
        q_idx = cols.index("question") if "question" in cols else 0
        a_idx = cols.index("answer")   if "answer" in cols and len(cols) > 1 else (1 if len(cols) > 1 else None)
        out = pd.DataFrame()
        out["question"] = df.iloc[:, q_idx].astype(str)
        out["answer"]   = df.iloc[:, a_idx].astype(str) if a_idx is not None else ""
        return out

    def _choose_name_col(self, df: pd.DataFrame) -> Optional[str]:
        if df.empty: return None
        prefs = ["product_name","description","item_name","product","name","title","sku_description"]
        for p in prefs:
            for c in df.columns:
                if p in c.lower():
                    return c
        return None

    def _clean_series_of_names(self, s: pd.Series) -> pd.Series:
        return (
            s.astype(str)
             .str.replace("\u00A0"," ", regex=False)
             .str.replace(r"\s+", " ", regex=True)
             .str.strip()
        )

    def _build_index(self):
        names = set()

        sku = self.tables.get("sku_master", pd.DataFrame())
        if not sku.empty:
            name_col = self._choose_name_col(sku)
            if name_col:
                names.update(self._clean_series_of_names(sku[name_col]).dropna())

        sales = self.tables.get("sales_transactions", pd.DataFrame())
        if not sales.empty:
            desc_col = None
            for c in sales.columns:
                if any(k in c.lower() for k in ["desc","name","product","item"]):
                    desc_col = c; break
            if not desc_col:
                desc_col = self._choose_name_col(sales)
            if desc_col:
                names.update(self._clean_series_of_names(sales[desc_col]).dropna())

        cat = self.tables.get("catalogue_flat", pd.DataFrame())
        if not cat.empty:
            name_col = self._choose_name_col(cat) or (list(cat.columns)[0] if len(cat.columns) else None)
            if name_col:
                names.update(self._clean_series_of_names(cat[name_col]).dropna())

        self.product_index = sorted({n for n in names if isinstance(n, str) and len(n) > 2})
        log.info(f"[index] size={len(self.product_index)}; sample={self.product_index[:5]}")

    def _load(self):
        log.info(f"[load] data_dir: {self.dir}")

        # Catalogue (DOCX → optional CSV)
        cat = self._find_file(FILE_HINTS["catalogue"], (".docx",))
        self.catalogue = self._parse_catalogue_docx(cat)
        if not self.catalogue.empty:
            log.info(f"[load] catalogue (docx): {len(self.catalogue)} raw lines")
            self._flatten_catalogue_to_csv()

        # Tabular sources (CSV / XLSX)
        for key, hints in FILE_HINTS.items():
            if key == "catalogue": 
                continue
            p = self._find_file(hints, (".csv",".xlsx"))
            if not p:
                self.tables[key] = pd.DataFrame()
                log.info(f"[load] {key}: 0 rows (file not found)")
                continue
            df = self._safe_csv(p) if p.suffix.lower()==".csv" else self._safe_excel(p)
            if key == "faq" and not df.empty:
                df = self._normalize_faq(df)
            self.tables[key] = df
            log.info(f"[load] {key}: {len(df)} rows; columns: {list(df.columns)[:20]}")

        # Ingest Catalogue_Flat.csv if present
        cat_flat_path = self.dir / "Catalogue_Flat.csv"
        if cat_flat_path.exists():
            cat_flat = self._safe_csv(cat_flat_path)
            if not cat_flat.empty:
                self.tables["catalogue_flat"] = cat_flat
                log.info(f"[load] catalogue_flat: {len(cat_flat)} rows; columns: {list(cat_flat.columns)[:20]}")

        self._build_index()

# -----------------------------
# Customer QA (data-only, English-core; multilingual via route wrapper)
# -----------------------------
class CustomerQA:
    def __init__(self, dm: DataManager):
        self.dm = dm
        self.synonyms = {
            "cashew nuts":"cashew", "cashews":"cashew", "peanuts":"peanut",
            "almonds":"almond", "apricots":"apricot", "pistachios":"pistachio",
            "unsalted":"unsalted", "baked":"baked", "healthy":"healthy"
        }
        self.healthy_tokens = ["baked", "unsalted", "fruit", "berries", "almond", "pistachio", "natural"]

    def _normalize(self, s: str) -> str:
        s = unidecode(str(s)).lower().strip()
        for k,v in self.synonyms.items():
            s = re.sub(rf"\b{k}\b", v, s, flags=re.I)
        s = re.sub(r"\s+", " ", s)
        return s

    def smalltalk(self, q: str) -> str:
        ql = q.lower()
        if any(w in ql for w in ["thanks","thank you"]): return "You’re welcome! If you need anything else, I’m here."
        if any(w in ql for w in ["bye","goodbye"]):      return "Bye for now! Have a great day."
        if any(w in ql for w in ["how are you","how r u","how’s it going","hows it going"]):
            return "I’m doing great and ready to help with snacks!"
        return FRIENDLY_GREETS[0]

    # --- FAQ matching: answer-only (never echo a question) ---
    def faq_best_match(self, query: str) -> tuple[int, str, str]:
        """Return (score, matched_question, matched_answer). Score is 0–100."""
        faq = self.dm.tables.get("faq", pd.DataFrame())
        if faq.empty:
            return (0, "", "")
        qn = self._normalize(query)
        best_score, best_q, best_a = -1, "", ""
        for _, r in faq.iterrows():
            q_txt = str(r.get("question","")).strip()
            a_txt = str(r.get("answer","")).strip()
            cand  = self._normalize(q_txt)
            score = max(fuzz.WRatio(qn, cand), fuzz.token_set_ratio(qn, cand), fuzz.partial_ratio(qn, cand))
            if score > best_score:
                best_score, best_q, best_a = score, q_txt, a_txt
        return (max(best_score,0), best_q, best_a)

    def answer_from_faq(self, query: str) -> Optional[str]:
        score, q_txt, a_txt = self.faq_best_match(query)
        if score >= 70 and a_txt:
            return a_txt
        return None

    def product_matches(self, query: str, limit=12) -> List[str]:
        if not self.dm.product_index: return []
        q = self._normalize(query)

        def _score(a, b):
            return max(
                fuzz.WRatio(a, b),
                fuzz.token_set_ratio(a, b),
                fuzz.partial_ratio(a, b),
            )

        scored = [(name, _score(q, name)) for name in self.dm.product_index]
        scored.sort(key=lambda x: x[1], reverse=True)
        hits = [name for name, s in scored if s >= 55]
        return hits[:limit]

    def healthy_suggestions(self, limit=5) -> List[str]:
        if not self.dm.product_index: return []
        picks = []
        for token in self.healthy_tokens:
            for m in self.product_matches(token, limit=limit*2):
                if m not in picks:
                    picks.append(m)
                if len(picks) >= limit:
                    return picks
        return picks

    def answer_en(self, raw_en: str) -> str:
        q = self._normalize(raw_en)

        # Order placement
        if re.search(r"\b(order|place.*order|how.*order)\b", q):
            a = self.answer_from_faq("order")
            return a if a else f"{FRIENDLY_SORRY[0]} I couldn’t find order instructions in the FAQ CSV."

        # Shipping / delivery / ETA / tracking
        if re.search(r"\b(ship|shipping|deliver|delivery|arrive|arrival|eta|status|track|tracking)\b", q):
            a = self.answer_from_faq("shipping") or self.answer_from_faq("delivery") or self.answer_from_faq("tracking")
            return a if a else f"{FRIENDLY_SORRY[0]} I couldn’t find shipping/ETA info in the FAQ CSV."

        # Refunds / returns
        if re.search(r"\b(refund|return|replace)\b", q):
            a = self.answer_from_faq("refund") or self.answer_from_faq("return")
            return a if a else f"{FRIENDLY_SORRY[0]} I couldn’t find refund/return info in the FAQ CSV."

        # Allergens / nutrition
        if re.search(r"\b(allergen|ingredient|nutrition|peanut|nut)\b", q):
            a = self.answer_from_faq(raw_en) or self.answer_from_faq("allergen") or self.answer_from_faq("ingredient")
            if a:
                return a
            hits = self.product_matches(raw_en)
            if hits:
                return f"{FRIENDLY_OK[0]} Related products:\n" + "\n".join(f"- {h}" for h in hits) + "\nTell me which one and I’ll check details."
            return f"{FRIENDLY_SORRY[0]} I couldn’t find allergen or nutrition details in the CSVs."

        # Healthy suggestions
        if re.search(r"(healthy|diet|light|low cal|low-calorie|low calorie)", q):
            hits = self.healthy_suggestions()
            if hits:
                return "Here are some lighter picks:\n" + "\n".join(f"- {h}" for h in hits) + "\nWant more like these?"
            return f"{FRIENDLY_SORRY[0]} I couldn’t find healthier options in the product list."

        # Generic product search (and short “have … ?” messages)
        if re.search(r"(do you have|have you got|looking for|i want|i need|find|search|any)\b", q) or len(q.split()) <= 6:
            hits = self.product_matches(raw_en)
            if hits:
                return f"{FRIENDLY_OK[0]} I found these:\n" + "\n".join(f"- {h}" for h in hits) + "\nTell me which one you mean and I can check availability."

        # FAQ fallback (answer only; never echo a question)
        a = self.answer_from_faq(raw_en)
        if a:
            return a

        return f"{FRIENDLY_SORRY[0]} I couldn’t find that in the FAQ or products. If you share a product name or order number, I’ll check again."

# -----------------------------
# Staff analytics (English core; translation wrapper in route)
# -----------------------------
class StaffAnalytics:
    def __init__(self, dm: DataManager):
        self.dm = dm

    @staticmethod
    def _norm(s: str) -> str:
        return re.sub(r'[^a-z0-9]+', '', str(s).lower())

    def _pick_col(self, df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
        if df.empty: return None
        norm_map = {c: self._norm(c) for c in df.columns}
        for want in candidates:
            w = self._norm(want)
            for c, nc in norm_map.items():
                if w == nc or w in nc:
                    return c
        return None

    def _parse_plan(self, text: str) -> Dict[str, Any]:
        q = text.strip()
        start, end = parse_date_range(q)
        op, n = "top", 5
        m = re.search(r"\btop\s+(\d+)", q, re.I)
        if m: op, n = "top", int(m.group(1))
        m = re.search(r"\bbottom\s+(\d+)", q, re.I)
        if m: op, n = "bottom", int(m.group(1))
        if re.search(r"\bmedian\b", q, re.I): op = "median"
        if re.search(r"\b(average|avg|mean)\b", q, re.I): op = "average"
        m_kw = re.search(r"\b(?:of|for)\s+([A-Za-z][\w\s\-\&/]+)$", q, re.I)
        kw = m_kw.group(1).strip() if m_kw else None

        # Only show SKUs if explicitly requested
        wants_sku = bool(re.search(r"\b(sku|product code|item code|item number|item no)\b", q, re.I))
        return {"operation":op,"n":n,"start":start,"end":end,"keyword":kw,"wants_sku":wants_sku}

    def sales(self, text_en: str) -> str:
        plan = self._parse_plan(text_en)
        df = self.dm.tables.get("sales_transactions", pd.DataFrame())
        if df.empty: return "I couldn’t find sales records yet."
        df = filter_by_date(df, plan["start"], plan["end"])

        revenue_cols = ["Total_Amount_Incl_VAT","Total_Amount_Excl_VAT","Net_Amount_Excl_VAT","Line_Amount","Amount","Revenue"]
        rev_col = next((c for c in revenue_cols if c in df.columns), None)
        if not rev_col: return "Sales file is missing a revenue amount column."
        df = df.copy()
        df[rev_col] = pd.to_numeric(df[rev_col], errors="coerce").fillna(0)

        # Prefer Description (product names) directly from Sales
        desc_col = self._pick_col(df, ["Description","Desc","Item Description","Product Description"])

        # Identify possible SKU/code in sales
        sku_in_sales = self._pick_col(df, ["SKU","Item No","Item Number","Item Code","Item","Product Code","Product_Code","SKU_Number"])

        # Identify SKU Master columns if present
        sku_df = self.dm.tables.get("sku_master", pd.DataFrame())
        sku_code = self._pick_col(sku_df, ["SKU","Item No","Item Number","Item Code","Product Code","Code"]) if not sku_df.empty else None
        sku_name = self._pick_col(sku_df, ["Product_Name","Product Name","Name","Description","Title"]) if not sku_df.empty else None

        kw = plan.get("keyword")
        wants_sku = plan.get("wants_sku", False)

        # CASE A: Use product names in Sales (Description) by default
        if desc_col and not wants_sku:
            work = df
            if kw:
                work = work[work[desc_col].astype(str).str.contains(kw, case=False, na=False)]
                if work.empty: return f"No matching sales for '{kw}' in that period."
            grouped = work.groupby(desc_col)[rev_col].sum().sort_values(ascending=False)
            label_fmt = lambda k, v: f"- {k}: ${v:,.2f}"

        # CASE B: Join Sales SKU -> SKU Master Product_Name (still names by default)
        elif sku_in_sales and sku_df is not None and not sku_df.empty and sku_code and sku_name and not wants_sku:
            slim_sales = df[[sku_in_sales, rev_col]].copy()
            agg_sales = slim_sales.groupby(sku_in_sales, dropna=False)[rev_col].sum().reset_index()
            join_df = agg_sales.merge(
                sku_df[[sku_code, sku_name]].drop_duplicates(),
                left_on=sku_in_sales, right_on=sku_code, how="left"
            )
            name_col = sku_name
            if kw:
                join_df = join_df[join_df[name_col].astype(str).str.contains(kw, case=False, na=False)]
                if join_df.empty: return f"No matching sales for '{kw}' in that period."
            grouped = join_df.groupby(name_col, dropna=False)[rev_col].sum().sort_values(ascending=False)
            label_fmt = lambda k, v: f"- {k}: ${v:,.2f}"

        # CASE C: User explicitly asked for SKUs
        elif wants_sku and sku_in_sales:
            work = df
            key_col = sku_in_sales
            if kw:
                if desc_col:
                    work = work[work[desc_col].astype(str).str.contains(kw, case=False, na=False)]
                else:
                    work = work[work[key_col].astype(str).str.contains(kw, case=False, na=False)]
                if work.empty: return f"No matching sales for '{kw}' in that period."
            grouped = work.groupby(key_col)[rev_col].sum().sort_values(ascending=False)
            label_fmt = lambda k, v: f"- {k}: ${v:,.2f}"

        # CASE D: Fallback to any name-like column in sales
        else:
            name_fallback = None
            for c in df.columns:
                if any(k in c.lower() for k in ["name","title","product"]):
                    name_fallback = c; break
            if not name_fallback:
                return "Sales file is missing product names (no Description and no usable SKU→name mapping)."
            work = df
            if kw:
                work = work[work[name_fallback].astype(str).str.contains(kw, case=False, na=False)]
                if work.empty: return f"No matching sales for '{kw}' in that period."
            grouped = work.groupby(name_fallback)[rev_col].sum().sort_values(ascending=False)
            label_fmt = lambda k, v: f"- {k}: ${v:,.2f}"

        if grouped.empty:
            return "No sales found in that period."

        op, n = plan["operation"], plan["n"]
        if op == "top":
            head = grouped.head(n)
            return "\n".join([f"Top {len(head)} products by revenue:"] + [label_fmt(k, v) for k, v in head.items()])
        if op == "bottom":
            tail = grouped.sort_values(ascending=True).head(n)
            return "\n".join([f"Bottom {len(tail)} products by revenue:"] + [label_fmt(k, v) for k, v in tail.items()])
        if op == "median":
            return f"Median revenue per product: ${grouped.median():,.2f}"
        if op == "average":
            suffix = f" for {kw}" if kw else ""
            return f"Average revenue per product{suffix}: ${grouped.mean():,.2f}"
        return f"Total revenue: ${grouped.sum():,.2f}"

    def customers(self) -> str:
        df = self.dm.tables.get("customers", pd.DataFrame())
        if df.empty: return "I couldn’t find customer records yet."
        lines = [f"Total customers: {len(df)}"]
        if "Segment" in df.columns:
            for k,v in df["Segment"].value_counts().items():
                lines.append(f"- {k}: {v}")
        if "LTV" in df.columns:
            try:
                avg = pd.to_numeric(df["LTV"], errors="coerce").mean()
                lines.append(f"Average LTV: ${avg:,.2f}")
            except: pass
        return "\n".join(lines)

    def traffic(self) -> str:
        df = self.dm.tables.get("ecommerce_purchases", pd.DataFrame())
        if df.empty: return "I couldn’t find ecommerce traffic records yet."
        if "Platform" in df.columns and "Revenue" in df.columns:
            sums = df.groupby("Platform")["Revenue"].sum().sort_values(ascending=False)
            return "\n".join([f"- {k}: ${v:,.2f}" for k,v in sums.items()])
        return "Traffic table loaded, but it doesn’t have Platform/Revenue columns."

    def events(self) -> str:
        df = self.dm.tables.get("events", pd.DataFrame())
        if df.empty: return "I couldn’t find event records yet."
        lines=[f"Total events: {len(df)}"]
        if "Attendance" in df.columns:
            try:
                avg = pd.to_numeric(df["Attendance"], errors="coerce").mean()
                lines.append(f"Average attendance: {avg:.0f}")
            except: pass
            name_col = next((c for c in df.columns if "name" in c.lower() or "event" in c.lower()), None)
            if name_col:
                totals = df.groupby(name_col)["Attendance"].sum().sort_values(ascending=False)
                if len(totals):
                    lines.append(f"Most popular event: {totals.index[0]} ({int(totals.iloc[0])} total attendees)")
        return "\n".join(lines)

# -----------------------------
# Flask app + Frontend (Template-safe HTML)
# -----------------------------
app = Flask(__name__)

INDEX_HTML_TPL = Template("""
<!doctype html>
<html>
<head>
  <title>${CUSTOMER}</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
  :root{ --gold:#d2b48c; --gold-dark:#c19a6b; --bg:#f7f5f2; }
  *{box-sizing:border-box}
  body{background:var(--bg);font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial,sans-serif;margin:0}
  .container{max-width:720px;margin:0 auto;padding:20px}
  .widget{background:#fff;border-radius:18px;box-shadow:0 8px 30px rgba(0,0,0,.08);overflow:hidden}
  .header{background:linear-gradient(135deg,var(--gold),#e9c77d);color:#2b2b2b;padding:20px;text-align:center}
  .header h1{margin:0 0 6px;font-size:22px}
  .header p{margin:0;opacity:.85;font-size:13px}
  .chat{height:520px;overflow-y:auto;padding:18px;display:flex;flex-direction:column;gap:12px;background:#fbfaf7}
  .bubble{padding:12px 14px;border-radius:16px;max-width:85%;line-height:1.4;white-space:pre-wrap;position:relative}
  .user{align-self:flex-end;background:var(--gold);color:#1f160d}
  .bot{align-self:flex-start;background:#ffffff;color:#2b2b2b;border:1px solid #eee}
  .badge{font-size:11px;font-weight:600;opacity:.8;margin-bottom:4px}
  .composer{display:flex;padding:14px;border-top:1px solid #eee;background:#fff}
  .input-wrap{flex:1;position:relative}
  textarea{width:100%;border-radius:24px;padding:12px 88px 12px 14px;border:2px solid #eee;resize:none;font-family:inherit;font-size:14px;outline:none}
  textarea:focus{border-color:var(--gold)}
  button{position:absolute;right:8px;top:50%;transform:translateY(-50%);background:var(--gold);color:#1f160d;border:none;padding:8px 16px;border-radius:18px;cursor:pointer;font-size:14px}
  button:hover{background:var(--gold-dark)}
  .typing{color:#666;font-style:italic;align-self:flex-start}
  .hint{font-size:12px;color:#6b6b6b;margin-top:6px}
  </style>
</head>
<body>
  <div class="container">
    <div class="widget">
      <div class="header">
        <h1>${CUSTOMER}</h1>
        <p>${CUSTOMER} serves customers from CSV data only. ${STAFF} answers staff analytics. Prefix with “Staff query:”</p>
      </div>
      <div id="chat" class="chat">
        <div class="bubble bot">
          <div class="badge">${CUSTOMER}</div>
Hello! I'm ${CUSTOMER}. I can help you with:
• Product information and availability
• Orders and shipping (FAQ-based)
• Nutritional info and allergens (FAQ or catalogue-based)
• For staff analytics, start your message with “Staff query:”
        </div>
        <div class="hint">Tip: You can chat in your language. Staff analytics still need the “Staff query:” prefix.</div>
      </div>
      <div class="composer">
        <div class="input-wrap">
          <textarea id="msg" rows="1" placeholder="Type your message…" onkeydown="handleEnter(event)"></textarea>
          <button onclick="sendMsg()">Send</button>
        </div>
      </div>
    </div>
  </div>

<script>
function handleEnter(e){ if(e.key==='Enter' && !e.shiftKey){ e.preventDefault(); sendMsg(); } }
function autoResize(el){ el.style.height='auto'; el.style.height=el.scrollHeight + 'px'; }
document.getElementById('msg').addEventListener('input', function(){ autoResize(this); });

function escapeHtml(str){
  return str.replace(/[&<>'"]/g, function(tag){
    const chars = { '&': '&amp;', '<': '&lt;', '>': '&gt;', "'": '&#39;', '"': '&quot;' };
    return chars[tag] || tag;
  });
}

async function sendMsg(){
  const box=document.getElementById('msg');
  const text=box.value.trim();
  if(!text) return;
  const chat=document.getElementById('chat');
  chat.insertAdjacentHTML('beforeend','<div class="bubble user">'+escapeHtml(text)+'</div>');
  box.value=''; autoResize(box); chat.scrollTop=chat.scrollHeight;

  const typing=document.createElement('div');
  typing.className='typing'; typing.innerText='Assistant is typing…';
  chat.appendChild(typing); chat.scrollTop=chat.scrollHeight;

  let resp = await fetch('/chat', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({message:text})});
  let data = await resp.json();
  chat.removeChild(typing);

  const speaker = escapeHtml(data.speaker || '${CUSTOMER}');
  const reply   = escapeHtml(data.reply || '');
  chat.insertAdjacentHTML('beforeend','<div class="bubble bot"><div class="badge">'+speaker+'</div>'+reply+'</div>');
  chat.scrollTop=chat.scrollHeight;
}
</script>
</body>
</html>
""")

# -----------------------------
# Instantiate core components
# -----------------------------
dm     = DataManager(DATA_DIR)
cust   = CustomerQA(dm)
staff  = StaffAnalytics(dm)
app    = Flask(__name__)

# -----------------------------
# Routes
# -----------------------------
@app.route("/")
def index():
    html = INDEX_HTML_TPL.safe_substitute(
        CUSTOMER=ASSISTANT_NAME_CUSTOMER,
        STAFF=ASSISTANT_NAME_STAFF
    )
    return render_template_string(html)

@app.route("/chat", methods=["POST"])
def chat():
    raw = (request.get_json(force=True).get("message") or "").strip()
    if not raw:
        return jsonify({"reply":"Please enter a message.","speaker":ASSISTANT_NAME_CUSTOMER})

    user_lang = detect_lang_safe(raw)

    # ===== Staff analytics (requires literal prefix "Staff query:") =====
    if raw.lower().startswith("staff query:"):
        staff_text = raw.split(":",1)[1].strip() if ":" in raw else ""
        staff_text_en = to_english_keywords(staff_text, user_lang) if needs_translation(user_lang) else staff_text
        q = staff_text_en.lower()

        if any(k in q for k in ["campaign","roi"]):
            reply_en = "Campaign analytics not implemented in this build."
        elif any(k in q for k in ["customer","segment","ltv"]):
            reply_en = staff.customers()
        elif any(k in q for k in ["traffic","platform","channel"]):
            reply_en = staff.traffic()
        elif any(k in q for k in ["event","attendance"]):
            reply_en = staff.events()
        else:
            reply_en = staff.sales(staff_text_en)

        reply = from_english_plain(reply_en, user_lang) if needs_translation(user_lang) else reply_en
        return jsonify({"reply": reply, "speaker": ASSISTANT_NAME_STAFF})

    # ===== Customer mode (data-only answers) =====
    if is_smalltalk(raw):
        reply_en = cust.smalltalk(raw)
        reply = from_english_plain(reply_en, user_lang) if needs_translation(user_lang) else reply_en
        return jsonify({"reply": reply, "speaker": ASSISTANT_NAME_CUSTOMER})

    query_en = to_english_keywords(raw, user_lang) if needs_translation(user_lang) else raw
    reply_en = cust.answer_en(query_en)
    reply    = from_english_plain(reply_en, user_lang) if needs_translation(user_lang) else reply_en
    return jsonify({"reply": reply, "speaker": ASSISTANT_NAME_CUSTOMER})

# -----------------------------
# Debug endpoints
# -----------------------------
@app.route("/debug/sources", methods=["GET"])
def debug_sources():
    info = {}
    for key, df in dm.tables.items():
        info[key] = {
            "rows": int(len(df)),
            "columns": list(map(str, df.columns[:30])),
        }
    return jsonify({
        "data_dir": str(dm.dir),
        "loaded_tables": info,
        "product_index_size": len(dm.product_index),
        "product_index_sample": dm.product_index[:15],
    })

@app.route("/debug/search")
def debug_search():
    term = (request.args.get("term") or "").strip()
    if not term:
        return jsonify({"error":"provide ?term=..."}), 400

    contains = [n for n in dm.product_index if term.lower() in n.lower()][:25]

    def _score(a, b):
        return max(
            fuzz.WRatio(a, b),
            fuzz.token_set_ratio(a, b),
            fuzz.partial_ratio(a, b),
        )
    scored = sorted([(n, _score(term.lower(), n.lower())) for n in dm.product_index],
                    key=lambda x: x[1], reverse=True)[:15]

    return jsonify({
        "term": term,
        "contains_hits": contains,
        "fuzzy_top": [{"name": n, "score": s} for n, s in scored],
    })

@app.route("/health")
def health():
    return jsonify({"ok": True})

# -----------------------------
# Entrypoint
# -----------------------------
if __name__ == "__main__":
    print("🚀 Starting server on http://127.0.0.1:5000/")
    app.run(host="127.0.0.1", port=5000, debug=False)
